import unittest
from unittest.mock import MagicMock, patch
import io
import openpyxl
import docx

from app.file_parser import parse_file

class FileParserTests(unittest.TestCase):
    def test_parse_csv(self):
        csv_content = b"col1,col2\nval1,val2\n"
        result = parse_file(csv_content, "data.csv")
        self.assertIn("col1, col2", result)
        self.assertIn("val1, val2", result)

    def test_parse_csv_utf16_with_nul_bytes(self):
        # UTF-16 CSV file containing NUL bytes when read as raw bytes
        csv_utf16 = "col1,col2\nval1,val2\n".encode("utf-16")
        result = parse_file(csv_utf16, "data.csv")
        self.assertNotIn("\x00", result)
        self.assertIn("col1, col2", result)
        self.assertIn("val1, val2", result)

    def test_parse_csv_semicolon_delimiter(self):
        # CSV with semicolon delimiter (common in Spanish locales)
        csv_semicolon = b"col1;col2\nval1;val2\n"
        result = parse_file(csv_semicolon, "data.csv")
        self.assertIn("col1, col2", result)
        self.assertIn("val1, val2", result)

    def test_parse_text_utf16(self):
        # UTF-16 text file
        text_utf16 = "Hello world\nLine 2".encode("utf-16")
        result = parse_file(text_utf16, "info.txt")
        self.assertEqual(result, "Hello world\nLine 2")

    def test_parse_text(self):
        text_content = "Hello world\nLine 2".encode("utf-8")
        result = parse_file(text_content, "info.txt")
        self.assertEqual(result, "Hello world\nLine 2")

    def test_parse_markdown(self):
        md_content = "# Title\nSome content".encode("utf-8")
        result = parse_file(md_content, "doc.md")
        self.assertEqual(result, "# Title\nSome content")

    def test_parse_docx(self):
        doc = docx.Document()
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Paragraph 2")
        out = io.BytesIO()
        doc.save(out)
        docx_bytes = out.getvalue()

        result = parse_file(docx_bytes, "test.docx")
        self.assertIn("Paragraph 1", result)
        self.assertIn("Paragraph 2", result)

    def test_parse_xlsx(self):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "Name"
        ws["B1"] = "Age"
        ws["A2"] = "Alice"
        ws["B2"] = 30
        out = io.BytesIO()
        wb.save(out)
        xlsx_bytes = out.getvalue()

        result = parse_file(xlsx_bytes, "test.xlsx")
        self.assertIn("--- Hoja: Sheet1 ---", result)
        self.assertIn("Name, Age", result)
        self.assertIn("Alice, 30", result)

    @patch("pypdf.PdfReader")
    def test_parse_pdf(self, mock_reader_class):
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "PDF Text Content"
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        mock_reader_class.return_value = mock_reader

        result = parse_file(b"dummy_pdf_bytes", "test.pdf")
        self.assertEqual(result, "PDF Text Content")
        mock_reader_class.assert_called_once()

    def test_unsupported_format(self):
        with self.assertRaises(ValueError) as context:
            parse_file(b"bytes", "image.png")
        self.assertIn("Formato de archivo no soportado", str(context.exception))

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "manual_assets"
OUT = ROOT / "docs" / "Manual_Administrador_Visual_Asistto.docx"

INK = "151716"
MUTED = "68706C"
GREEN = "176B5B"
LIGHT_GREEN = "E8F1ED"
BG = "F4F5F2"
LINE = "DDE2DC"
PANEL = "FFFFFF"
AMBER = "AD6500"
BLUE = "315F9F"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def rounded(draw, xy, fill, outline=None, width=1, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def text(draw, xy, value, size=28, fill="#" + INK, bold=False, anchor=None):
    draw.text(xy, value, font=font(size, bold=bold), fill=fill, anchor=anchor)


def pill(draw, x, y, label, fill="#E8F1ED", color="#" + GREEN):
    w = max(104, len(label) * 11 + 32)
    rounded(draw, (x, y, x + w, y + 34), fill=fill, radius=17)
    text(draw, (x + 16, y + 8), label, size=16, fill=color, bold=True)
    return w


def screen_base(title, subtitle):
    img = Image.new("RGB", (1600, 900), "#" + BG)
    d = ImageDraw.Draw(img)
    rounded(d, (70, 70, 1530, 830), fill="#FBFCFA", outline="#" + LINE, width=2, radius=26)
    rounded(d, (70, 70, 325, 830), fill="#FFFFFF", outline="#" + LINE, width=2, radius=26)
    rounded(d, (105, 110, 155, 160), fill="#" + INK, radius=10)
    text(d, (119, 124), "WA", size=19, fill="#FFFFFF", bold=True)
    text(d, (170, 112), "WhatsApp Bot", size=24, bold=True)
    text(d, (170, 142), "Panel admin", size=16, fill="#" + MUTED)
    for i, label in enumerate(["Dashboard", "Clientes", "Bots", "Conversaciones", "CRM"]):
        y = 210 + i * 66
        fill = "#E8EEE9" if label in title or (title == "Mapa del admin" and label == "Dashboard") else "#FFFFFF"
        rounded(d, (105, y, 290, y + 46), fill=fill, radius=12)
        text(d, (132, y + 13), label, size=17, fill="#" + INK)
    text(d, (380, 120), title, size=42, bold=True)
    text(d, (382, 174), subtitle, size=22, fill="#" + MUTED)
    return img, d


def save_map():
    img, d = screen_base("Mapa del admin", "Rutas principales para operar la plataforma")
    nodes = [
        ("Login", 430, 290, GREEN),
        ("Clientes", 650, 220, BLUE),
        ("Bots", 910, 220, BLUE),
        ("Prompt", 1160, 220, GREEN),
        ("Knowledge", 1160, 360, GREEN),
        ("Conversaciones", 650, 500, AMBER),
        ("CRM", 910, 500, AMBER),
        ("Diagnosticos", 1160, 500, MUTED),
    ]
    for label, x, y, color in nodes:
        rounded(d, (x, y, x + 185, y + 86), fill="#FFFFFF", outline="#" + color, width=3, radius=18)
        text(d, (x + 24, y + 29), label, size=22, fill="#" + color, bold=True)
    lines = [
        ((615, 333), (650, 263)), ((835, 263), (910, 263)),
        ((1095, 263), (1160, 263)), ((1095, 263), (1160, 403)),
        ((615, 333), (650, 543)), ((835, 543), (910, 543)),
        ((1095, 543), (1160, 543)),
    ]
    for a, b in lines:
        d.line((a, b), fill="#" + LINE, width=5)
    img.save(ASSETS / "01-mapa-admin.png")


def save_client_flow():
    img, d = screen_base("Clientes", "Alta de cliente, usuarios y primer bot")
    steps = [
        ("1", "Crear cliente", "Nombre y slug unico"),
        ("2", "Crear usuario", "client_admin o client_viewer"),
        ("3", "Crear bot", "Nombre, slug y Phone Number ID"),
        ("4", "Configurar agente", "Prompt y base de conocimiento"),
    ]
    for i, (num, title, subtitle) in enumerate(steps):
        x = 410 + i * 265
        rounded(d, (x, 300, x + 230, 520), fill="#FFFFFF", outline="#" + LINE, width=2, radius=22)
        rounded(d, (x + 22, 326, x + 72, 376), fill="#" + GREEN, radius=25)
        text(d, (x + 39, 337), num, size=24, fill="#FFFFFF", bold=True)
        text(d, (x + 22, 405), title, size=24, bold=True)
        text(d, (x + 22, 446), subtitle, size=17, fill="#" + MUTED)
        if i < len(steps) - 1:
            d.line((x + 232, 410, x + 260, 410), fill="#" + GREEN, width=5)
    rounded(d, (420, 610, 1400, 700), fill="#" + LIGHT_GREEN, outline="#BFD7CE", width=2, radius=18)
    text(d, (452, 635), "Resultado", size=21, bold=True, fill="#" + GREEN)
    text(d, (452, 668), "El cliente ya puede entrar al panel y operar su bot con permisos propios.", size=20, fill="#" + INK)
    img.save(ASSETS / "02-alta-cliente.png")


def save_bot_config():
    img, d = screen_base("Bots", "Configuracion operativa de un bot")
    rounded(d, (410, 250, 1450, 720), fill="#FFFFFF", outline="#" + LINE, width=2, radius=24)
    text(d, (450, 292), "Bot Clinica Demo", size=34, bold=True)
    pill(d, 450, 344, "active")
    fields = [
        ("Cliente", "Clinica Demo"),
        ("Phone Number ID", "2158310464922380"),
        ("Prompt", "Publicado en Postgres"),
        ("Knowledge", "3 documentos activos"),
        ("Prueba", "WhatsApp real + conversaciones"),
    ]
    for i, (label, value) in enumerate(fields):
        y = 425 + i * 50
        text(d, (455, y), label, size=19, fill="#" + MUTED, bold=True)
        text(d, (705, y), value, size=20, fill="#" + INK)
    rounded(d, (1120, 300, 1390, 365), fill="#" + INK, radius=14)
    text(d, (1152, 319), "Editar prompt", size=22, fill="#FFFFFF", bold=True)
    rounded(d, (1120, 390, 1390, 455), fill="#FFFFFF", outline="#" + INK, width=2, radius=14)
    text(d, (1152, 409), "Base de conocimiento", size=21, fill="#" + INK, bold=True)
    img.save(ASSETS / "03-configurar-bot.png")


def save_prompt_editor():
    img, d = screen_base("Prompt", "Editor del comportamiento del agente")
    rounded(d, (410, 245, 1450, 735), fill="#FFFFFF", outline="#" + LINE, width=2, radius=22)
    text(d, (450, 285), "Instrucciones del agente", size=24, bold=True)
    prompt_lines = [
        "Eres el asistente de WhatsApp de Clinica Demo.",
        "Tu objetivo es resolver dudas, calificar interesados y agendar citas.",
        "",
        "Reglas:",
        "- Responde breve, claro y amable.",
        "- No inventes precios.",
        "- Si el usuario quiere una cita, pide nombre, motivo, dia y hora.",
        "- Si no sabes algo, ofrece pasar el caso a una persona.",
    ]
    rounded(d, (450, 330, 1395, 640), fill="#FAFBF9", outline="#CBD3CC", width=2, radius=14)
    for i, line in enumerate(prompt_lines):
        text(d, (480, 358 + i * 33), line, size=21 if i < 2 else 18, fill="#" + INK)
    rounded(d, (450, 665, 650, 720), fill="#" + GREEN, radius=12)
    text(d, (478, 682), "Publicar prompt", size=20, fill="#FFFFFF", bold=True)
    img.save(ASSETS / "04-editor-prompt.png")


def save_knowledge():
    img, d = screen_base("Base de conocimiento", "Documentos que alimentan al bot")
    docs = [
        ("Servicios", "Limpieza dental, blanqueamiento, ortodoncia..."),
        ("Horarios", "Lunes a viernes de 9:00 a 18:00"),
        ("Politicas", "Para precios exactos, ofrecer cita de valoracion"),
    ]
    for i, (title, body) in enumerate(docs):
        y = 260 + i * 130
        rounded(d, (410, y, 1020, y + 96), fill="#FFFFFF", outline="#" + LINE, width=2, radius=16)
        text(d, (440, y + 22), title, size=24, bold=True)
        text(d, (440, y + 58), body, size=18, fill="#" + MUTED)
        pill(d, 850, y + 30, "active")
    rounded(d, (1080, 260, 1450, 655), fill="#FFFFFF", outline="#" + LINE, width=2, radius=18)
    text(d, (1115, 300), "Agregar documento", size=25, bold=True)
    text(d, (1115, 365), "Titulo", size=18, fill="#" + MUTED)
    rounded(d, (1115, 395, 1410, 440), fill="#FAFBF9", outline="#CBD3CC", radius=10)
    text(d, (1135, 407), "FAQ", size=17, fill="#" + MUTED)
    text(d, (1115, 475), "Contenido", size=18, fill="#" + MUTED)
    rounded(d, (1115, 505, 1410, 590), fill="#FAFBF9", outline="#CBD3CC", radius=10)
    rounded(d, (1115, 615, 1350, 670), fill="#" + GREEN, radius=12)
    text(d, (1144, 632), "Guardar documento", size=19, fill="#FFFFFF", bold=True)
    img.save(ASSETS / "05-knowledge.png")


def save_roles():
    img, d = screen_base("Roles y permisos", "Que puede hacer cada usuario")
    headers = ["Accion", "Agencia", "Cliente admin", "Cliente viewer"]
    rows = [
        ("Crear clientes", "Si", "No", "No"),
        ("Crear bots", "Si", "No", "No"),
        ("Editar prompt", "Si", "Si", "No"),
        ("Editar knowledge", "Si", "Si", "No"),
        ("Ver conversaciones", "Si", "Si", "Si"),
        ("Ver CRM", "Si", "Si", "Si"),
    ]
    x0, y0 = 405, 250
    widths = [400, 190, 240, 240]
    y = y0
    for i, h in enumerate(headers):
        x = x0 + sum(widths[:i])
        rounded(d, (x, y, x + widths[i], y + 56), fill="#E8EEF5", outline="#" + LINE, radius=8)
        text(d, (x + 18, y + 17), h, size=19, bold=True, fill="#" + BLUE)
    for r, row in enumerate(rows):
        y = y0 + 62 + r * 62
        for i, val in enumerate(row):
            x = x0 + sum(widths[:i])
            rounded(d, (x, y, x + widths[i], y + 56), fill="#FFFFFF", outline="#" + LINE, radius=8)
            color = "#" + GREEN if val == "Si" else "#9B1C1C" if val == "No" else "#" + INK
            label = "✓" if val == "Si" else "—" if val == "No" else val
            text(d, (x + 18, y + 16), label, size=22, bold=val in ("Si", "No"), fill=color)
    img.save(ASSETS / "06-roles.png")


def save_daily_ops():
    img, d = screen_base("Operacion diaria", "Rutina para mantener el bot saludable")
    items = [
        ("Conversaciones", "Confirmar mensajes y respuestas"),
        ("CRM", "Revisar leads calificados"),
        ("Escalaciones", "Atender casos humanos"),
        ("IA", "Probar modelo y prompt"),
        ("Calendario", "Validar agenda"),
        ("Mejoras", "Ajustar knowledge"),
    ]
    for i, (title, body) in enumerate(items):
        x = 410 + (i % 3) * 335
        y = 260 + (i // 3) * 200
        rounded(d, (x, y, x + 300, y + 145), fill="#FFFFFF", outline="#" + LINE, width=2, radius=20)
        rounded(d, (x + 24, y + 24, x + 72, y + 72), fill="#" + LIGHT_GREEN, radius=24)
        text(d, (x + 39, y + 35), str(i + 1), size=20, fill="#" + GREEN, bold=True)
        text(d, (x + 92, y + 28), title, size=23, bold=True)
        text(d, (x + 24, y + 92), body, size=18, fill="#" + MUTED)
    img.save(ASSETS / "07-operacion-diaria.png")


def add_run(p, value, size=11, bold=False, color=INK):
    r = p.add_run(value)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return r


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, value, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, value, size=10, bold=bold, color=color)


def add_h1(doc, value):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    add_run(p, value, size=16, bold=True, color="2E74B5")


def add_h2(doc, value):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    add_run(p, value, size=13, bold=True, color="2E74B5")


def add_body(doc, value, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    add_run(p, value, size=11, bold=bold)


def add_bullets(doc, values):
    for value in values:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        add_run(p, value, size=10.5)


def add_image(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(ASSETS / filename), width=Inches(6.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    add_run(cap, caption, size=9.5, color=MUTED)


def build_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)
    save_map()
    save_client_flow()
    save_bot_config()
    save_prompt_editor()
    save_knowledge()
    save_roles()
    save_daily_ops()


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "Asistto | Manual visual del administrador", size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Documento operativo interno", size=9, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(50)
    title.paragraph_format.space_after = Pt(8)
    add_run(title, "Manual Visual del Administrador", size=25, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "WhatsApp Bot Multi-Cliente | Asistto", size=14, color=MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(24)
    add_run(meta, "Guia para alta de clientes, configuracion de bots, prompt, conocimiento y operacion diaria", size=10.5, color=MUTED)
    add_image(doc, "01-mapa-admin.png", "Vista general de las rutas principales del panel.")

    doc.add_page_break()
    add_h1(doc, "1. Como se organiza el panel")
    add_body(doc, "El panel esta pensado para operar muchos clientes desde una sola instalacion. La agencia administra todo; cada cliente entra solo a sus bots.")
    add_image(doc, "06-roles.png", "Matriz de permisos por rol.")
    add_bullets(doc, [
        "Agencia: crea clientes, bots y usuarios, y edita cualquier agente.",
        "Cliente admin: edita prompt y base de conocimiento de sus bots.",
        "Cliente viewer: consulta conversaciones, CRM y estado sin editar configuracion.",
    ])

    add_h1(doc, "2. Alta de cliente")
    add_body(doc, "El alta inicia en /admin/clients. Primero se crea el cliente, despues su usuario y finalmente su primer bot.")
    add_image(doc, "02-alta-cliente.png", "Flujo recomendado para dar de alta un cliente.")
    add_bullets(doc, [
        "Usa slugs cortos, en minusculas y sin espacios.",
        "Crea al menos un usuario client_admin por cliente.",
        "Entrega la contrasena temporal por un canal seguro.",
    ])

    add_h1(doc, "3. Configuracion del bot")
    add_body(doc, "El detalle del bot concentra el estado operativo y los accesos a la configuracion del agente.")
    add_image(doc, "03-configurar-bot.png", "Pantalla conceptual de configuracion del bot.")
    add_bullets(doc, [
        "Phone Number ID debe venir de Meta WhatsApp Cloud API.",
        "El numero visible solo ayuda a identificar el bot en el panel.",
        "Los mensajes entrantes se enrutan por phone_number_id.",
    ])

    doc.add_page_break()
    add_h1(doc, "4. Prompt del agente")
    add_body(doc, "El prompt define identidad, tono, reglas comerciales, preguntas clave, limites y comportamiento de agenda o escalacion.")
    add_image(doc, "04-editor-prompt.png", "Editor visual del prompt publicado.")
    add_bullets(doc, [
        "Guardar el prompt publica una version activa en Postgres.",
        "No requiere redeploy para reflejarse en nuevas respuestas.",
        "Si un bot no tiene prompt propio, usa el fallback versionado en prompts/system.md.",
    ])

    add_h1(doc, "5. Base de conocimiento")
    add_body(doc, "La base de conocimiento complementa el prompt con informacion factual del cliente: servicios, horarios, precios, politicas y preguntas frecuentes.")
    add_image(doc, "05-knowledge.png", "Pantalla conceptual de documentos de conocimiento.")
    add_bullets(doc, [
        "Los documentos activos se agregan al prompt runtime del bot.",
        "Los documentos archivados no se usan para responder.",
        "Mantener documentos cortos y concretos mejora la calidad de respuesta.",
    ])

    add_h1(doc, "6. Operacion diaria")
    add_body(doc, "La rutina diaria se enfoca en confirmar que WhatsApp recibe mensajes, que la IA responde, que el CRM avanza y que las escalaciones no se queden sin atender.")
    add_image(doc, "07-operacion-diaria.png", "Checklist visual de operacion diaria.")

    doc.add_page_break()
    add_h1(doc, "7. Rutas principales")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Ruta", "Uso", "Acceso"]
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=BLUE)
    routes = [
        ("/admin/login", "Entrada al panel", "Todos"),
        ("/admin/clients", "Crear y revisar clientes", "Agencia"),
        ("/admin/clients/{client_id}", "Crear bots y usuarios cliente", "Agencia"),
        ("/admin/bots", "Listado de bots", "Agencia y clientes"),
        ("/admin/bots/{bot_id}", "Detalle operativo del bot", "Agencia y clientes"),
        ("/admin/bots/{bot_id}/prompt", "Editar prompt activo", "Agencia, client_admin"),
        ("/admin/bots/{bot_id}/knowledge", "Crear/listar conocimiento", "Agencia, client_admin"),
        ("/admin/conversations", "Revisar conversaciones", "Agencia y clientes"),
        ("/admin/crm", "Gestionar leads", "Agencia y clientes"),
        ("/admin/escalations", "Atender escalaciones", "Agencia y clientes"),
        ("/admin/ai-status", "Diagnostico IA", "Agencia"),
        ("/admin/calendar-status", "Diagnostico calendario", "Agencia"),
        ("/admin/reset-contact", "Reset de contactos de prueba", "Agencia"),
    ]
    for route, use, access in routes:
        cells = table.add_row().cells
        set_cell_text(cells[0], route)
        set_cell_text(cells[1], use)
        set_cell_text(cells[2], access)

    add_h1(doc, "8. Checklists")
    add_h2(doc, "Alta de cliente")
    add_bullets(doc, [
        "Crear cliente.",
        "Crear usuario client_admin.",
        "Crear bot.",
        "Registrar Phone Number ID.",
        "Configurar prompt.",
        "Cargar knowledge.",
        "Probar WhatsApp real.",
        "Confirmar conversacion en admin.",
    ])
    add_h2(doc, "Configuracion del bot")
    add_bullets(doc, [
        "Nombre claro y slug unico.",
        "Phone Number ID correcto.",
        "Prompt publicado.",
        "Knowledge activo.",
        "Tono validado por el cliente.",
        "Prueba real completada.",
    ])

    doc.add_page_break()
    add_h1(doc, "9. Guia para laminas con gpt-image-2")
    add_body(doc, "Los graficos de este documento usan un estilo visual preparado para reinterpretarse con gpt-image-2 si se desea producir una version mas editorial o comercial.")
    prompts = [
        ("Mapa general", "Infografia SaaS en espanol con rutas Login, Dashboard, Clientes, Bots, Prompt, Knowledge, Conversaciones, CRM, Diagnosticos."),
        ("Alta de cliente", "Flujo numerado en espanol: crear cliente, crear usuario, crear bot, configurar agente, probar WhatsApp."),
        ("Roles", "Matriz visual de permisos para Agencia, Cliente admin y Cliente viewer."),
        ("Configuracion de bot", "Pantalla de admin para bot con Phone Number ID, prompt, knowledge y prueba de WhatsApp."),
        ("Operacion diaria", "Checklist visual de operacion diaria para conversaciones, CRM, escalaciones, IA, calendario y mejoras."),
    ]
    for title_text, prompt_text in prompts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_run(p, title_text + ": ", bold=True, color=GREEN)
        add_run(p, prompt_text, color=INK)

    doc.save(OUT)


if __name__ == "__main__":
    build_assets()
    build_docx()
    print(OUT)
